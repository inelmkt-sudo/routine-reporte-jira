# -*- coding: utf-8 -*-
"""Genera el reporte diario Jira en Word a partir de data/report.json.

Uso: python scripts/render_docx.py data/report.json
Salida: reports/reporte-jira-YYYY-MM-DD.docx
"""
import json
import os
import sys

from docx import Document
from docx.shared import Pt, RGBColor

AZUL = RGBColor(0x1F, 0x4E, 0x79)
ROJO = RGBColor(0xC0, 0x00, 0x00)


def fix_mojibake(s):
    """Corrige textos UTF-8 mal decodificados como latin-1 (Ã³ -> ó)."""
    if not isinstance(s, str):
        return s
    if "Ã" in s or "Â" in s:
        try:
            return s.encode("latin-1").decode("utf-8")
        except (UnicodeEncodeError, UnicodeDecodeError):
            return s
    return s


def fix_all(obj):
    if isinstance(obj, dict):
        return {k: fix_all(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [fix_all(v) for v in obj]
    return fix_mojibake(obj)


def main():
    if len(sys.argv) != 2:
        print("Uso: render_docx.py <report.json>", file=sys.stderr)
        sys.exit(1)

    with open(sys.argv[1], encoding="utf-8") as fh:
        data = fix_all(json.load(fh))

    doc = Document()
    doc.add_heading(f"Reporte diario Jira — {data['fecha']}", 0)
    doc.add_paragraph("INEL — Actividad de las últimas 24 horas, todos los proyectos.")

    # Resumen ejecutivo
    r = data["resumen"]
    doc.add_heading("Resumen ejecutivo", 1)
    t = doc.add_table(rows=1, cols=5)
    t.style = "Light Grid Accent 1"
    hdr = ["Issues tocados", "Creados", "Completados", "En progreso", "Bloqueados"]
    vals = [r.get("total_issues", 0), r.get("creados", 0), r.get("completados", 0),
            r.get("en_progreso", 0), r.get("bloqueados", 0)]
    for i, h in enumerate(hdr):
        t.rows[0].cells[i].text = h
    row = t.add_row()
    for i, v in enumerate(vals):
        row.cells[i].text = str(v)

    # Alertas
    alertas = r.get("alertas", [])
    doc.add_heading("Alertas de trazabilidad", 1)
    if alertas:
        for a in alertas:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(a)
            run.font.color.rgb = ROJO
    else:
        doc.add_paragraph("Sin alertas: no se detectaron fechas movidas, borrados ni tareas vencidas.")

    # Iniciativas transversales (análisis macro por palabras clave)
    doc.add_heading("Iniciativas transversales", 1)
    inis = data.get("iniciativas", [])
    if inis:
        t = doc.add_table(rows=1, cols=5)
        t.style = "Light Grid Accent 1"
        for i, h in enumerate(["Iniciativa", "Avance", "Pods", "Integrantes", "Issues (hoy)"]):
            t.rows[0].cells[i].text = h
        for ini in inis:
            row = t.add_row()
            row.cells[0].text = ini["clave"]
            av = ini.get("avance") or {}
            if av:
                row.cells[1].text = (f"{av.get('porcentaje', 0)}% "
                                     f"({av.get('completados', 0)}/{av.get('total', 0)} done, "
                                     f"{av.get('en_progreso', 0)} en curso, "
                                     f"{av.get('pendientes', 0)} pend.)")
            else:
                row.cells[1].text = "—"
            row.cells[2].text = ", ".join(ini.get("pods", []))
            row.cells[3].text = ", ".join(ini.get("integrantes", []))
            row.cells[4].text = ", ".join(ini.get("issues", []))
    else:
        doc.add_paragraph("No se detectaron iniciativas que crucen pods o integrantes en esta ventana.")

    # Detalle exhaustivo por integrante (tabla compacta)
    doc.add_heading("Actividad por integrante (detalle completo)", 1)
    for persona in data.get("integrantes", []):
        h = doc.add_heading(f"{persona['nombre']} — {len(persona['issues'])} issue(s)", 2)
        for run in h.runs:
            run.font.color.rgb = AZUL
        if persona.get("total_horas_ciclo"):
            p = doc.add_paragraph()
            run = p.add_run(f"Horas de ciclo del día: {persona['total_horas_ciclo']}")
            run.bold = True
        t = doc.add_table(rows=1, cols=7)
        t.style = "Light Grid Accent 1"
        for i, hdr in enumerate(["Issue", "Resumen", "Estado", "Acción", "Horas ciclo", "Vence", "Pertenece a"]):
            t.rows[0].cells[i].text = hdr
        for iss in persona["issues"]:
            row = t.add_row()
            row.cells[0].text = f"{iss['key']}\n({iss['proyecto']})"
            row.cells[1].text = iss["resumen"]
            row.cells[2].text = iss["estado"]
            row.cells[3].text = iss["accion"]
            row.cells[4].text = iss.get("horas_ciclo") or "—"
            row.cells[5].text = iss.get("vencimiento") or "—"
            row.cells[6].text = iss.get("padre") or "—"
        for cell in t.rows[0].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
        # Cambios y comentarios debajo de la tabla, solo para issues que los tengan
        for iss in persona["issues"]:
            extras = [f"Cambio: {c}" for c in iss.get("cambios", [])]
            extras += [f"Comentario: {c}" for c in iss.get("comentarios", [])]
            for e in extras:
                p = doc.add_paragraph(f"{iss['key']} — {e}", style="List Bullet")
                p.paragraph_format.left_indent = Pt(18)
                for run in p.runs:
                    run.font.size = Pt(9)

    os.makedirs("reports", exist_ok=True)
    fecha_archivo = data.get("fecha_archivo") or data["fecha"].split()[0]
    fecha_archivo = fecha_archivo.replace("/", "-")
    out = os.path.join("reports", f"reporte-jira-{fecha_archivo}.docx")
    doc.save(out)
    print(json.dumps({"ok": True, "file": out}))


if __name__ == "__main__":
    main()
